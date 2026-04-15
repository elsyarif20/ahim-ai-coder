import streamlit as st
from groq import Groq
from supabase import create_client, Client
import base64
import uuid
from datetime import datetime
from docx import Document
from fpdf import FPDF
import io
import re
import hashlib
import pandas as pd
from gtts import gTTS
from scipy import stats
import statsmodels.api as sm
from duckduckgo_search import DDGS
import xlsxwriter

# --- 1. HELPERS & PROFESSIONAL EXPORT ENGINE ---
def search_web(query):
    """Pencarian Real-time di Backend"""
    try:
        with DDGS() as ddgs:
            results = [r for r in ddgs.text(query, max_results=5)]
            return "\n".join([f"Sumber: {r['title']} - {r['body']}" for r in results]) if results else ""
    except: return ""

def parse_markdown_table(text):
    """Parsing Markdown ke DataFrame"""
    try:
        lines = [l.strip() for l in text.split('\n') if '|' in l]
        if len(lines) < 3: return None
        header = [c.strip() for c in lines[0].split('|') if c.strip()]
        data = []
        for l in lines[2:]:
            row = [c.strip() for c in l.split('|') if c.strip()]
            if len(row) == len(header): data.append(row)
        return pd.DataFrame(data, columns=header)
    except: return None

def export_to_excel(df):
    """Export ke Excel dengan Border & Header Profesional"""
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    df.to_excel(writer, index=False, sheet_name='Ahim_Research')
    workbook  = writer.book
    worksheet = writer.sheets['Ahim_Research']
    
    # Format Border & Header
    fmt_border = workbook.add_format({'border': 1, 'align': 'left'})
    fmt_header = workbook.add_format({'bold': True, 'border': 1, 'bg_color': '#D7E4BC', 'align': 'center'})
    
    for col_num, value in enumerate(df.columns.values):
        worksheet.write(0, col_num, value, fmt_header)
        worksheet.set_column(col_num, col_num, 25)
    
    for row_num, row_data in enumerate(df.values):
        for col_num, cell_data in enumerate(row_data):
            worksheet.write(row_num + 1, col_num, cell_data, fmt_border)
            
    writer.close()
    return output.getvalue()

def export_to_docx(text):
    """Export ke Word Profesional Tanpa Simbol Markdown"""
    doc = Document()
    doc.add_heading('Laporan Penelitian Ahim Devpad', 0)
    clean_txt = re.sub(r'[\*\#\`]', '', text)
    for para in clean_txt.split('\n'):
        if para.strip(): doc.add_paragraph(para)
    bio = io.BytesIO(); doc.save(bio)
    return bio.getvalue()

def export_to_pdf(text):
    """Export ke PDF dengan Margin Bersih"""
    pdf = FPDF()
    pdf.add_page(); pdf.set_font("Arial", size=11)
    clean_txt = re.sub(r'[\*\#\`]', '', text)
    pdf.multi_cell(0, 8, txt=clean_txt, align='L')
    return bytes(pdf.output())

def speak_text(text):
    try:
        clean_txt = re.sub(r'[\*\#\`]', '', text)
        tts = gTTS(text=clean_txt, lang='id')
        fp = io.BytesIO(); tts.write_to_fp(fp)
        return fp.getvalue()
    except: return None

# --- 2. UI & AUTH ---
st.set_page_config(page_title="Ahim Devpad v6.8", layout="wide")
url, key = st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"]
supabase = create_client(url, key)

if "user" not in st.session_state: st.session_state.user = None
if not st.session_state.user:
    c1, c2, c3 = st.columns([1, 1.5, 1])
    with c2:
        st.markdown("<h1 style='text-align:center;'>AHIM DEVPAD</h1>", unsafe_allow_html=True)
        t1, t2 = st.tabs(["🔐 Login", "📝 Daftar"])
        with t1:
            with st.form("l"):
                e, p = st.text_input("Email"), st.text_input("Password", type="password")
                if st.form_submit_button("Masuk Sekarang", use_container_width=True, type="primary"):
                    res = supabase.auth.sign_in_with_password({"email": e, "password": p})
                    if res.user: st.session_state.user = res.user; st.rerun()
    st.stop()

# --- 3. DATA LOADING ---
db_data = supabase.table("user_data").select("*").eq("email", st.session_state.user.email).single().execute().data
user_api_key = db_data.get("groq_api_key", "")
chat_history = db_data.get("chat_history", {}) or {}

if "current_chat_id" not in st.session_state:
    st.session_state.current_chat_id = list(chat_history.keys())[0] if chat_history else str(uuid.uuid4())
cur_id = st.session_state.current_chat_id
if cur_id not in chat_history:
    chat_history[cur_id] = {"title": "Baru", "messages": [], "timestamp": str(datetime.now())}

# --- 4. SIDEBAR ---
with st.sidebar:
    st.title("⚙️ Setelan")
    if st.button("🚪 Keluar"): 
        supabase.auth.sign_out(); st.session_state.user = None; st.rerun()
    with st.expander("🔑 API & Info"):
        new_api = st.text_input("Groq Key", value=user_api_key, type="password")
        if st.button("Simpan"):
            supabase.table("user_data").update({"groq_api_key": new_api}).eq("email", st.session_state.user.email).execute(); st.rerun()
        st.info("Ahim Devpad v6.8\nLiyas Syarifudin, M.Pd.\n2026")
    st.divider()
    if st.button("➕ Chat Baru"): st.session_state.current_chat_id = str(uuid.uuid4()); st.rerun()
    for cid, cdata in sorted(chat_history.items(), key=lambda x: x[1].get('timestamp',''), reverse=True):
        if st.button(f"💬 {cdata.get('title','')[:20]}", key=cid, use_container_width=True):
            st.session_state.current_chat_id = cid; st.rerun()

# --- 5. CHAT AREA ---
current_chat = chat_history.get(cur_id, {"messages": []})
for m in current_chat["messages"]:
    with st.chat_message(m["role"]):
        st.markdown(m["content"])
        if m["role"] == "assistant":
            h = hashlib.md5(m["content"].encode()).hexdigest()[:8]
            c1, c2, c3, c4, c5 = st.columns([1,1,1,1,2])
            c1.download_button("📄 TXT", m["content"].encode(), f"{h}.txt", key=f"t{h}")
            c2.download_button("🟦 Word", export_to_docx(m["content"]), f"{h}.docx", key=f"w{h}")
            c3.download_button("🟥 PDF", export_to_pdf(m["content"]), f"{h}.pdf", key=f"p{h}")
            if c4.button("🔊 Suara", key=f"v{h}"):
                audio = speak_text(m["content"])
                if audio: st.audio(audio, format='audio/mp3')
            df = parse_markdown_table(m["content"])
            if df is not None:
                c5.download_button("🟢 Copy to Sheet", export_to_excel(df), f"data_{h}.xlsx", key=f"x{h}")

# --- 6. INPUT ---
prompt = st.chat_input("Cari, diskusikan, atau tulis 1000 kata...", accept_audio=True)
if prompt:
    client = Groq(api_key=user_api_key)
    u_text = prompt.text if not prompt.audio else client.audio.transcriptions.create(file=("i.wav", prompt.audio.read()), model="whisper-large-v3-turbo", response_format="text")
    
    if u_text:
        is_long = any(x in u_text.lower() for x in ["1000 kata", "panjang", "detail"])
        needs_search = any(x in u_text.lower() for x in ["cari", "google", "referensi"])
        
        web_data = search_web(u_text) if needs_search else ""
        sys_msg = "Kamu Ahim AI, asisten riset buatan Liyas Syarifudin, M.Pd. Gunakan bahasa manusiawi."
        final_p = f"Referensi Web: {web_data}\n\n{u_text}" if web_data else u_text
        if is_long: final_p += " (Tulis minimal 1000 kata mendetail)."

        chat_history[cur_id]["messages"].append({"role": "user", "content": u_text})
        with st.chat_message("user"): st.markdown(u_text)

        with st.chat_message("assistant"):
            ph, full = st.empty(), ""
            stream = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role":"system","content":sys_msg}, *[{"role":m["role"],"content":m["content"]} for m in chat_history[cur_id]["messages"][:-1]], {"role":"user","content":final_p}],
                temperature=0.7, max_tokens=8192, stream=True
            )
            for chunk in stream:
                if (content := chunk.choices[0].delta.content):
                    full += content; ph.markdown(full + "▌")
            ph.markdown(full)
        
        chat_history[cur_id]["messages"].append({"role": "assistant", "content": full})
        chat_history[cur_id]["timestamp"] = str(datetime.now())
        supabase.table("user_data").update({"chat_history": chat_history}).eq("email", st.session_state.user.email).execute()
        st.rerun()