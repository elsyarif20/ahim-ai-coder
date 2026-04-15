import io, re, pandas as pd, hashlib
from duckduckgo_search import DDGS
from docx import Document
from fpdf import FPDF
from gtts import gTTS
import xlsxwriter
import PyPDF2

def search_web(query):
    try:
        with DDGS() as ddgs:
            results = [r for r in ddgs.text(query, max_results=5)]
            return "\n".join([f"Sumber: {r['title']} - {r['body']}" for r in results]) if results else ""
    except: return ""

def parse_markdown_table(text):
    try:
        lines = [l.strip() for l in text.split('\n') if '|' in l]
        if len(lines) < 3: return None
        header = [c.strip() for c in lines[0].split('|') if c.strip()]
        data = [[c.strip() for c in l.split('|') if c.strip()] for l in lines[2:]]
        return pd.DataFrame([d for d in data if len(d) == len(header)], columns=header)
    except: return None

def export_excel_pro(df):
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    df.to_excel(writer, index=False, sheet_name='Ahim_Research')
    workbook, worksheet = writer.book, writer.sheets['Ahim_Research']
    fmt_border = workbook.add_format({'border': 1, 'align': 'left'})
    fmt_header = workbook.add_format({'bold': True, 'border': 1, 'bg_color': '#D7E4BC', 'align': 'center'})
    for col_num, value in enumerate(df.columns.values):
        worksheet.write(0, col_num, value, fmt_header)
        worksheet.set_column(col_num, col_num, 20)
    for row_num, row_data in enumerate(df.values):
        for col_num, cell_data in enumerate(row_data):
            worksheet.write(row_num + 1, col_num, cell_data, fmt_border)
    writer.close()
    return output.getvalue()

def export_docx_pro(text):
    doc = Document()
    doc.add_heading('Laporan Penelitian Ahim Devpad', 0)
    clean = re.sub(r'[\*\#\`]', '', text)
    for p in clean.split('\n'):
        if p.strip(): doc.add_paragraph(p)
    bio = io.BytesIO(); doc.save(bio)
    return bio.getvalue()

def export_pdf_pro(text):
    pdf = FPDF()
    pdf.add_page(); pdf.set_font("Arial", size=11)
    clean = re.sub(r'[\*\#\`]', '', text)
    pdf.multi_cell(0, 8, txt=clean, align='L')
    return bytes(pdf.output())

def speak_text(text):
    try:
        clean = re.sub(r'[\*\#\`]', '', text)
        tts = gTTS(text=clean, lang='id')
        fp = io.BytesIO(); tts.write_to_fp(fp)
        return fp.getvalue()
    except: return None
